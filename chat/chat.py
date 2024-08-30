import os
import requests
from openai import OpenAI
import time
import tiktoken
import openai
import streamlit as st
from docx import Document
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter
from langchain.document_loaders import DirectoryLoader
from langchain.embeddings import OpenAIEmbeddings, HuggingFaceInstructEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain import OpenAI
from openai import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.memory import ConversationBufferMemory
from tqdm.autonotebook import tqdm
from langchain.chains import ConversationalRetrievalChain
from html_copy import css, bot_template, user_template


load_dotenv()
client = OpenAI(
    api_key=os.environ.get('OPENAI_API_KEY')
)


def llm(text, system_message, delimiter="####", print_response=False, retries=3, sleep_time=10):
    while retries > 0:
        # Define messages with the user's input and a system message
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": f"{delimiter}{text}{delimiter}"}
        ]

        # Calculate max_tokens for output/completion
        encoding = tiktoken.encoding_for_model("gpt-3.5-turbo-16k-0613")
        max_tokens = 16000 - (len(encoding.encode(text)) + len(encoding.encode(system_message))) - 13
        # Call to LLM model
        try:
            completion = client.chat.completions.create(
                model="gpt-3.5-turbo-16k-0613",
                messages=messages,
                max_tokens=max_tokens,
                temperature=0
            )
            response = completion.choices[0].message.content
            if print_response:
                print(response)
            break
        except openai.RateLimitError as e:
            print('Catching RateLimitError, retrying in 1 minute...')
            retries -= 1
            time.sleep(sleep_time)

    return response



def get_job_text(documents):
    text_list = []  # Create a list to store text from all documents
    for doc in documents:
        file_extension = doc.name.split(".")[-1].lower()

        if file_extension == "pdf":
            pdf_reader = PdfReader(doc)
            for page in pdf_reader.pages:
                text_list.append(page.extract_text())

        elif file_extension in ["docx", "doc"]:
            docx = Document(doc)
            for paragraph in docx.paragraphs:
                text_list.append(paragraph.text)
                
    job_descrip = "\n".join(text_list)
      
    return job_descrip

# Reading CV's and summarizing the content using LLM
def get_pdf_text(documents_cv, document_job):
    candidates_score = []

    for doc in documents_cv:
        text_list = []  # Create a list to store text from the current document
        file_extension = doc.name.split(".")[-1].lower()

        if file_extension == "pdf":
            pdf_reader = PdfReader(doc)
            for page in pdf_reader.pages:
                text_list.append(page.extract_text())

        elif file_extension in ["docx", "doc"]:
            docx = Document(doc)
            for paragraph in docx.paragraphs:
                text_list.append(paragraph.text)

        combined_text = "\n".join(text_list)
        system_message_compare = f" I have a job ads {document_job}\
            I have  cv from a candidates/applicants applying for the position.\
                        please provide a significant reasonings why the candidate is suitable for the job.\
                            Please note, every paragraph should have the candidates name."
        llm_score = llm(combined_text, system_message_compare, delimiter="####", print_response=True, retries=3, sleep_time=10)
        candidates_score.append(llm_score)

    candidates_score_text = '\n'.join(candidates_score)
    return candidates_score_text


def get_text_chunks(text):
    text_splitter =CharacterTextSplitter(chunk_size = 1000,
                                         chunk_overlap = 50)
    chunks = text_splitter.split_text(text)
    return chunks

def get_vectorstore(text_chunks):
    embeddings = OpenAIEmbeddings()
    vectorstore = FAISS.from_texts(texts=text_chunks, embedding=embeddings)
    return vectorstore

def get_conversation_chain(vectorstore):
    llm = ChatOpenAI(temperature=0, model = "gpt-3.5-turbo-16k-0613")
    memory = ConversationBufferMemory(
        memory_key = 'chat_history', return_messages = True)
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm = llm,
        retriever = vectorstore.as_retriever(search_kwargs={'fetch_k': 80,'k': 80}),
        memory = memory
    )
    return conversation_chain

def handle_userinput(user_question):
    response = st.session_state.conversation({'question': user_question})
    st.session_state.chat_history = response['chat_history']

    for i, message in enumerate(st.session_state.chat_history):
        if i % 2 == 0:
            st.write(user_template.replace("{{MSG}}", message.content), unsafe_allow_html=True)
        else:
            st.write(bot_template.replace("{{MSG}}", message.content), unsafe_allow_html=True)
            
def main():
    load_dotenv()
    st.set_page_config(page_title=" ", page_icon="")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.write(' ')
    with col2:
        st.image("cvbot.png",width=80)
    with col3:
        st.write(' ')

    st.write(css, unsafe_allow_html=True)
    st.markdown(css, unsafe_allow_html=True)

    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = None


    user_question = st.chat_input("Message CV MatchBot")
    #st.markdown("<h1 style='text-align: center; color: grey;'>Powered by OpenAI and streamlit</h1>", unsafe_allow_html=True)

    
    if user_question:
        handle_userinput(user_question)

    with st.sidebar:
        st.subheader("Your job ads")
        documents_job = st.file_uploader("Upload your job ads", accept_multiple_files=True)
        st.subheader("Your Curriculum Vitae's (CV's)")
        documents_cv = st.file_uploader("Upload your CV's", accept_multiple_files=True)
        if st.button("Process"):
            with st.spinner("Processing..."):
                job_text = get_job_text(documents_job)
                raw_text = get_pdf_text(documents_cv, job_text)
                text_chunks = get_text_chunks(raw_text)
                vectorstore = get_vectorstore(text_chunks)
                st.session_state.conversation = get_conversation_chain(vectorstore)

if __name__ == '__main__':
    main()


